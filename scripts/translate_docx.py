from __future__ import annotations

import argparse
from concurrent.futures import ThreadPoolExecutor, TimeoutError
import re
import time
from typing import Iterable

from docx import Document
from docx.text.paragraph import Paragraph
from deep_translator import GoogleTranslator


LETTER_RE = re.compile(r"[A-Za-zÁÉÍÓÖŐÚÜŰáéíóöőúüű]")


def looks_translatable(text: str) -> bool:
    if not text:
        return False
    normalized = text.strip()
    if not normalized:
        return False
    return bool(LETTER_RE.search(normalized))


def iter_all_paragraphs(doc: Document) -> Iterable[Paragraph]:
    for paragraph in doc.paragraphs:
        yield paragraph

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            yield paragraph
        for paragraph in section.footer.paragraphs:
            yield paragraph


def translate_texts(texts: list[str], source: str, target: str) -> list[str]:
    translated: list[str] = []

    def translate_one(text: str) -> str:
        translator = GoogleTranslator(source=source, target=target)
        return translator.translate(text)

    for i, text in enumerate(texts):
        for attempt in range(4):
            try:
                with ThreadPoolExecutor(max_workers=1) as pool:
                    future = pool.submit(translate_one, text)
                    translated.append(future.result(timeout=10))
                time.sleep(0.1)
                break
            except (TimeoutError, Exception):
                if attempt == 3:
                    translated.append(text)
                else:
                    time.sleep(0.8 + attempt * 0.4)
        if (i + 1) % 50 == 0:
            print(f"Translated {i + 1}/{len(texts)}")
    return translated


def main() -> None:
    parser = argparse.ArgumentParser(description="Translate DOCX text while keeping structure.")
    parser.add_argument("--input", required=True, help="Input DOCX path")
    parser.add_argument("--output", required=True, help="Output DOCX path")
    parser.add_argument("--source", default="hu", help="Source language code")
    parser.add_argument("--target", default="en", help="Target language code")
    args = parser.parse_args()

    doc = Document(args.input)
    paragraphs = list(iter_all_paragraphs(doc))

    original_texts: list[str] = []
    paragraph_indexes: list[int] = []
    for idx, paragraph in enumerate(paragraphs):
        raw = paragraph.text
        if looks_translatable(raw):
            original_texts.append(raw)
            paragraph_indexes.append(idx)

    translated_texts = translate_texts(original_texts, source=args.source, target=args.target)

    for para_idx, translated in zip(paragraph_indexes, translated_texts):
        paragraphs[para_idx].text = translated

    doc.save(args.output)
    print(f"Translated {len(translated_texts)} paragraphs to {args.output}")


if __name__ == "__main__":
    main()

